# ui/wizard.py
# Dataset Description Wizard (CICADAS-style) for TCIA submissions
#
# Drop this file at: TCIA-Data-Mine-Project/ui/wizard.py
# Then in app.py:
#   from ui.wizard import wizard_page
#   nav = st.sidebar.radio("Navigation", ["Submit", "Dataset Wizard", "Admin"], index=0)
#   if nav == "Dataset Wizard": wizard_page()
#
# This file is intentionally self-contained for a 1-week draft.
# You can later refactor call_ollama + prompts into common/ai.py.

from __future__ import annotations

import io
import json
import os
import shlex
import subprocess
from datetime import datetime
from typing import Dict, List, Optional, Tuple

import requests
import streamlit as st

OLLAMA_HOST = os.getenv("OLLAMA_HOST", "http://127.0.0.1:11434")

# ----------------------------
# Wizard section map (1-week draft)
# ----------------------------
SECTIONS: List[dict] = [
    {
        "key": "title",
        "label": "Title",
        "guidance": "A clear, descriptive dataset name. Avoid vague titles like 'Cancer Imaging Dataset'.",
        "placeholder": "Example: Multi-institutional CT imaging dataset for lung nodule detection (2018–2023)",
        "required": True,
    },
    {
        "key": "abstract",
        "label": "Abstract",
        "guidance": "2–5 sentences: what the dataset is, who/what it includes, modalities, and the intended purpose.",
        "placeholder": "Describe the dataset at a high level without deep method details.",
        "required": True,
    },
    {
        "key": "background",
        "label": "Background / Rationale",
        "guidance": "Why this dataset exists, the clinical/scientific motivation, and what gap it addresses.",
        "placeholder": "Explain the problem space and why sharing this dataset is useful.",
        "required": False,
    },
    {
        "key": "contents",
        "label": "Dataset Contents",
        "guidance": "What is included: DICOM, segmentations, labels, clinical tables, reports, code, etc.",
        "placeholder": "List the components and formats (DICOM, CSV, NIfTI, JSON, etc.).",
        "required": True,
    },
    {
        "key": "cohort",
        "label": "Cohort / Subjects",
        "guidance": "Who/what is represented, inclusion criteria, time range, and any key stratifications.",
        "placeholder": "Include inclusion/exclusion at a high level and what each subject represents.",
        "required": True,
    },
    {
        "key": "labels",
        "label": "Labels / Outcomes / Annotations",
        "guidance": "Describe ground truth: label definitions, who annotated, and what artifacts exist (segmentations, classes, etc.).",
        "placeholder": "Explain label meanings and how they were produced, without inventing details you do not have.",
        "required": False,
    },
]

REVIEW_KEY = "__review__"


# ----------------------------
# Ollama caller (copied from app.py style, simplified)
# ----------------------------
def call_ollama(
    model: str,
    prompt: str,
    temperature: float = 0.2,
    num_predict: int = 400,
    timeout_s: int = 120,
) -> str:
    """
    1) Try /api/chat (non-stream)
    2) Fallback to /api/generate (stream)
    3) Fallback to CLI
    Returns a readable error string instead of failing silently.
    """

    model = (model or "").strip()
    if not model:
        return "[Ollama error] Model name is empty."

    # 1) /api/chat
    try:
        resp = requests.post(
            f"{OLLAMA_HOST}/api/chat",
            json={
                "model": model,
                "messages": [{"role": "user", "content": prompt}],
                "options": {"temperature": temperature, "num_predict": num_predict},
                "stream": False,
            },
            timeout=timeout_s,
        )

        if resp.status_code == 200:
            data = resp.json()
            # Ollama chat usually returns {"message": {"role": "...", "content": "..."}, ...}
            if isinstance(data, dict):
                msg = data.get("message") or {}
                if isinstance(msg, dict) and "content" in msg:
                    return (msg.get("content") or "").strip()
                if "response" in data:
                    return (data.get("response") or "").strip()
            return "[Ollama error] Unexpected /api/chat response format."

        if resp.status_code != 404:
            return f"[Ollama error] /api/chat HTTP {resp.status_code}: {resp.text[:300]}"

    except Exception as e:
        chat_err = str(e)
    else:
        chat_err = None

    # 2) /api/generate streaming
    try:
        r = requests.post(
            f"{OLLAMA_HOST}/api/generate",
            json={
                "model": model,
                "prompt": prompt,
                "options": {"temperature": temperature, "num_predict": num_predict},
                "stream": True,
            },
            stream=True,
            timeout=timeout_s,
        )

        if r.status_code == 200:
            chunks = []
            for line in r.iter_lines(decode_unicode=True):
                if not line:
                    continue
                try:
                    obj = json.loads(line)
                except json.JSONDecodeError:
                    continue

                if obj.get("error"):
                    return f"[Ollama error] {obj['error']}"

                # generate endpoint emits "response" chunks
                if "response" in obj and obj["response"] is not None:
                    chunks.append(obj["response"])

                # stop if "done" appears
                if obj.get("done") is True:
                    break

            out = "".join(chunks).strip()
            return out or "[empty response]"

        if r.status_code != 404:
            return f"[Ollama error] /api/generate HTTP {r.status_code}: {r.text[:300]}"

    except Exception as e:
        gen_err = str(e)
    else:
        gen_err = None

    # 3) CLI fallback
    try:
        cmd = ["ollama", "run", model]
        result = subprocess.run(
            cmd,
            input=prompt,
            text=True,
            capture_output=True,
            timeout=180,
        )
        if result.returncode != 0:
            err = (result.stderr or "").strip()
            return "[Ollama error] " + (err or "CLI call failed.")
        return (result.stdout or "").strip()

    except Exception as e:
        cli_err = str(e)

    # If we got here, everything failed. Return the most useful error(s).
    bits = []
    if chat_err:
        bits.append(f"chat: {chat_err}")
    if gen_err:
        bits.append(f"generate: {gen_err}")
    bits.append(f"host: {OLLAMA_HOST}")
    return "[Ollama error] All methods failed. " + " | ".join(bits)

# ----------------------------
# Prompting
# ----------------------------
def feedback_prompt(section_label: str, section_text: str) -> str:
    section_text = (section_text or "").strip()
    return f"""
You are an expert technical editor helping an investigator write a dataset description in CICADAS style.

Section to improve: {section_label}

Rules:
- Do not invent specifics (counts, dates, scanner models, institutions) that are not present.
- If key info is missing, ask for it explicitly.
- Use clear professional language.
- Keep the response compact and actionable.

Output format (use these headings exactly):
Missing:
Suggestions:
Rewrite:

Text to review:
\"\"\"{section_text}\"\"\"
""".strip()


def compile_for_submit(answers: Dict[str, str]) -> str:
    """Compile a single text blob suitable for pre-filling the Submit form's dataset_description."""
    blocks = []
    for s in SECTIONS:
        k = s["key"]
        val = (answers.get(k) or "").strip()
        if not val:
            continue
        blocks.append(f"{s['label']}:\n{val}\n")
    return "\n".join(blocks).strip()


# ----------------------------
# DOCX export
# ----------------------------
def build_docx_bytes(
    answers: Dict[str, str],
    title_fallback: str = "Dataset Description",
) -> bytes:
    """
    Requires python-docx to be installed.
    """
    try:
        from docx import Document
    except Exception as e:
        raise RuntimeError("python-docx is required for Word export") from e

    doc = Document()

    title = (answers.get("title") or "").strip() or title_fallback
    doc.add_heading(title, level=0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    doc.add_paragraph("")

    for s in SECTIONS:
        k = s["key"]
        doc.add_heading(s["label"], level=1)
        val = (answers.get(k) or "").strip()
        if val:
            doc.add_paragraph(val)
        else:
            doc.add_paragraph("[TBD]")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ----------------------------
# Session state helpers
# ----------------------------
def _init_state():
    if "wizard_step" not in st.session_state:
        st.session_state.wizard_step = 0
    if "wizard_answers" not in st.session_state:
        st.session_state.wizard_answers = {s["key"]: "" for s in SECTIONS}
    if "wizard_feedback" not in st.session_state:
        st.session_state.wizard_feedback = {s["key"]: "" for s in SECTIONS}

    # NEW: dialog state
    if "wizard_show_feedback_dialog" not in st.session_state:
        st.session_state.wizard_show_feedback_dialog = False
    if "wizard_feedback_section_key" not in st.session_state:
        st.session_state.wizard_feedback_section_key = ""

    if "prefill" not in st.session_state:
        st.session_state.prefill = {
            "proposal_type": "new_collection",
            "email": "",
            "dataset_title": "",
            "dataset_description": "",
        }


def _current_total_steps() -> int:
    # sections + review step
    return len(SECTIONS) + 1


def _is_review_step(step_idx: int) -> bool:
    return step_idx >= len(SECTIONS)


# ----------------------------
# UI
# ----------------------------
def wizard_page():
    _init_state()

    if st.session_state.wizard_show_feedback_dialog:
        _feedback_dialog()

    st.title("Dataset Description Wizard")
    st.caption(
        "Step-by-step CICADAS-style drafting with local AI feedback. "
        "At the end you can export to Word and prefill the main Submit form."
    )

    # Sidebar helper panel (works alongside app.py sidebar nav)
    with st.sidebar:
        st.markdown("### Wizard Progress")
        total = _current_total_steps()
        step = int(st.session_state.wizard_step)
        st.write(f"Step {min(step + 1, total)} of {total}")

        labels = [s["label"] for s in SECTIONS] + ["Review & Export"]
        for i, lab in enumerate(labels):
            marker = "➡️ " if i == step else ""
            st.write(f"{marker}{lab}")

        st.markdown("---")
        col_a, col_b = st.columns(2)
        with col_a:
            if st.button("Reset wizard", use_container_width=True):
                st.session_state.wizard_step = 0
                st.session_state.wizard_answers = {s["key"]: "" for s in SECTIONS}
                st.session_state.wizard_feedback = {s["key"]: "" for s in SECTIONS}
                st.rerun()
        with col_b:
            if st.button("Load example", use_container_width=True):
                st.session_state.wizard_answers.update(
                    {
                        "title": "Example: Multi-site imaging dataset for [TBD] application",
                        "abstract": "This dataset contains [TBD] imaging studies collected from [TBD] between [TBD]. "
                                    "It is intended to support research in [TBD].",
                        "background": "The motivation for creating this dataset is to enable reproducible research in [TBD] "
                                      "and address limitations of existing datasets.",
                        "contents": "Included: DICOM images, a CSV with basic clinical variables, and a label file describing [TBD].",
                        "cohort": "The cohort includes subjects meeting [TBD] criteria. Exclusions include [TBD].",
                        "labels": "Labels include [TBD] and were produced by [TBD]. Definitions: [TBD].",
                    }
                )
                st.rerun()

    step = int(st.session_state.wizard_step)
    total = _current_total_steps()

    if _is_review_step(step):
        _render_review_and_export()
        return

    section = SECTIONS[step]
    key = section["key"]

    st.subheader(section["label"])
    if section.get("guidance"):
        st.info(section["guidance"])

    text_val = st.text_area(
        "Your draft",
        value=st.session_state.wizard_answers.get(key, ""),
        placeholder=section.get("placeholder", ""),
        height=220,
        key=f"wizard_input_{key}",
    )

    # Persist on each run
    st.session_state.wizard_answers[key] = text_val

    # AI feedback block
    st.markdown("#### AI feedback (optional)")
    model = st.text_input("Local model name", value="llama3.2:3b", key="wizard_model_name")

    with st.expander("Troubleshoot Ollama connection"):
        if st.button("Test connection"):
            try:
                t = requests.get(f"{OLLAMA_HOST}/api/tags", timeout=10)
                st.write("Status:", t.status_code)
                st.json(t.json() if t.status_code == 200 else {"body": t.text[:500]})
            except Exception as e:
                st.error(str(e))

    cols = st.columns([1, 1, 2])
    with cols[0]:
        run_ai = st.button("Get AI feedback", use_container_width=True)
    with cols[1]:
        clear_ai = st.button("Clear feedback", use_container_width=True)

    if clear_ai:
        st.session_state.wizard_feedback[key] = ""
        st.rerun()

    if run_ai:
        if not (text_val or "").strip():
            st.warning("Add some text first, then run feedback.")
        else:
            prompt = feedback_prompt(section["label"], text_val)
            with st.spinner("Running local AI feedback..."):
                out = call_ollama(model=model.strip(), prompt=prompt, temperature=0.2)
            st.session_state.wizard_feedback[key] = out

            # NEW: pop the feedback dialog immediately
            _open_feedback_dialog(key)
            st.rerun()


    fb = (st.session_state.wizard_feedback.get(key) or "").strip()
    if fb:
        st.text_area("Feedback", fb, height=240)

        # Simple helper: try to extract rewrite block
        rewrite = _extract_rewrite(fb)
        if rewrite:
            st.markdown("##### Suggested rewrite")
            st.text_area("Rewrite", rewrite, height=180, key=f"wizard_rewrite_{key}")

            col1, col2 = st.columns(2)
            with col1:
                if st.button("Replace my draft with rewrite", use_container_width=True):
                    st.session_state.wizard_answers[key] = rewrite
                    # Also update current text_area value on rerun
                    st.session_state[f"wizard_input_{key}"] = rewrite
                    st.rerun()
            with col2:
                if st.button("Append rewrite below my draft", use_container_width=True):
                    combined = (st.session_state.wizard_answers[key] or "").rstrip() + "\n\n" + rewrite
                    st.session_state.wizard_answers[key] = combined
                    st.session_state[f"wizard_input_{key}"] = combined
                    st.rerun()

    st.markdown("---")

    # Nav buttons
    back_disabled = step <= 0
    next_disabled = False
    if section.get("required") and not (st.session_state.wizard_answers.get(key) or "").strip():
        next_disabled = True

    nav_cols = st.columns([1, 1, 3])
    with nav_cols[0]:
        if st.button("Back", disabled=back_disabled, use_container_width=True):
            st.session_state.wizard_step = max(0, step - 1)
            st.rerun()
    with nav_cols[1]:
        if st.button("Next", disabled=next_disabled, use_container_width=True):
            st.session_state.wizard_step = min(total - 1, step + 1)
            st.rerun()
    with nav_cols[2]:
        if next_disabled:
            st.caption("This step is required before continuing.")

def _open_feedback_dialog(section_key: str):
    st.session_state.wizard_feedback_section_key = section_key
    st.session_state.wizard_show_feedback_dialog = True


@st.dialog("AI feedback")
def _feedback_dialog():
    key = st.session_state.wizard_feedback_section_key
    fb = (st.session_state.wizard_feedback.get(key) or "").strip()

    if not fb:
        st.info("No feedback to show.")
        if st.button("Close", use_container_width=True):
            st.session_state.wizard_show_feedback_dialog = False
            st.rerun()
        return

    st.text_area("Feedback", fb, height=260)

    rewrite = _extract_rewrite(fb)
    if rewrite:
        st.markdown("#### Suggested rewrite")
        st.text_area("Rewrite", rewrite, height=200, key=f"dialog_rewrite_{key}")

    c1, c2 = st.columns(2)

    with c1:
        # Accept: replace the user's draft with the rewrite if available,
        # otherwise do nothing except close.
        if st.button("Accept", type="primary", use_container_width=True):
            if rewrite:
                st.session_state.wizard_answers[key] = rewrite
                st.session_state[f"wizard_input_{key}"] = rewrite
            st.session_state.wizard_show_feedback_dialog = False
            st.rerun()

    with c2:
        # Decline: keep draft as-is, just close.
        if st.button("Decline", use_container_width=True):
            st.session_state.wizard_show_feedback_dialog = False
            st.rerun()

def _extract_rewrite(feedback_text: str) -> str:
    """
    Tries to parse the 'Rewrite:' portion out of the model output.
    If not present, returns empty string.
    """
    txt = feedback_text or ""
    lower = txt.lower()

    idx = lower.find("rewrite:")
    if idx < 0:
        return ""
    after = txt[idx + len("rewrite:") :].strip()

    # If the model repeats headings, stop at next heading
    stop_markers = ["missing:", "suggestions:"]
    stop_positions = []
    lower_after = after.lower()
    for m in stop_markers:
        j = lower_after.find(m)
        if j >= 0:
            stop_positions.append(j)
    if stop_positions:
        after = after[: min(stop_positions)].strip()

    return after.strip()


def _render_review_and_export():
    st.subheader("Review & Export")

    answers = st.session_state.wizard_answers

    st.caption("Review your sections below. You can go Back to edit, or export when ready.")
    for s in SECTIONS:
        k = s["key"]
        st.markdown(f"### {s['label']}")
        val = (answers.get(k) or "").strip()
        if val:
            st.write(val)
        else:
            st.warning("TBD")

    st.markdown("---")
    compiled = compile_for_submit(answers)
    st.markdown("### Combined text (for Submit page prefill)")
    st.text_area("Combined", compiled, height=220)

    c1, c2, c3 = st.columns([1, 1, 2])

    with c1:
        if st.button("Back", use_container_width=True):
            st.session_state.wizard_step = max(0, len(SECTIONS) - 1)
            st.rerun()

    with c2:
        if st.button("Prefill Submit page", use_container_width=True):
            title = (answers.get("title") or "").strip()
            st.session_state.prefill["dataset_title"] = title
            st.session_state.prefill["dataset_description"] = compiled
            st.success("Prefill set. Switch to Submit in the sidebar to see it populated.")

    with c3:
        try:
            docx_bytes = build_docx_bytes(answers, title_fallback="Dataset Description")
            safe_title = (answers.get("title") or "dataset_description").strip() or "dataset_description"
            filename = _safe_filename(safe_title) + ".docx"

            st.download_button(
                "Download Word document (.docx)",
                data=docx_bytes,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except Exception as e:
            st.error(f"Word export unavailable: {e}")


def _safe_filename(name: str) -> str:
    # Keep it simple and Windows-friendly.
    keep = []
    for ch in (name or ""):
        if ch.isalnum() or ch in (" ", "_", "-"):
            keep.append(ch)
    out = "".join(keep).strip().replace(" ", "_")
    return out[:80] if out else "dataset_description"

if __name__ == "__main__":
    wizard_page()

#Once the AI is done, it's output will be put in an extra box that shows up.
#From there, there is an accept or decline button that will either replace the text area with the AI output or leave it as is.